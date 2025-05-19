from docx import Document
from docx.shared import Pt

# 创建一个新的Word文档对象
doc = Document()

# 设置文档标题部分
title = doc.add_heading('', 0)  # 0表示标题级别为最高级（类似标题1）
title.paragraph_format.alignment = 1  # 1表示居中对齐
run = title.add_run("重 庆 理 工 大 学\n本科生课程目标达成情况评价及总结报告")
run.font.size = Pt(16)  # 设置字体大小

# 添加学年学期信息
year_semester = doc.add_paragraph()
year_semester.paragraph_format.alignment = 1
year_semester.add_run("2023 - 2024 学年第 1 学期").font.size = Pt(12)

# 留出一些空白行
for _ in range(10):
    doc.add_paragraph()

# 添加课程信息
course_info = doc.add_paragraph()
course_info.paragraph_format.alignment = 1
course_info.add_run("课程：程序设计基础(JAVA)I").font.size = Pt(12)

# 保存文档
doc.save("course_report.docx")