from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 创建文档
doc = Document()

# 1. 添加标题
title = doc.add_heading("年度销售报告", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
title.style.font.size = Pt(20)  # 字体大小
title.style.font.name = "黑体"  # 中文字体需指定

# 2. 添加段落
paragraph = doc.add_paragraph("以下是2024年各季度销售数据统计：")
paragraph.style.font.size = Pt(12)
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 3. 添加表格
data = [
    ["季度", "销售额（万元）", "增长率"],
    ["Q1", 500, "10%"],
    ["Q2", 600, "20%"],
    ["Q3", 550, "10%"],
    ["Q4", 700, "27%"]
]

table = doc.add_table(rows=len(data), cols=len(data[0]), style="Table Grid")
for i, row_data in enumerate(data):
    row = table.rows[i].cells
    for j, col_data in enumerate(row_data):
        row[j].text = str(col_data)
        # 设置表格内容居中对齐
        row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 4. 添加图片
# doc.add_picture("sales_chart.png", width=Inches(5))  # 插入图片并指定宽度

# 5. 保存文档
doc.save("sales_report.docx")