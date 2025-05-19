from docx import Document
import ast
import re
import logging

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s  %(levelname)s:%(message)s')
document = Document(r"C:\Users\lenovo\Desktop\分析报告.docx")
#第五个表格：试题与成绩分析表
table = document.tables[4]
with open(r"F:\代码\Python代码\课程达成度平台开发\pythonProject2\excel表统计数据",'r') as file:
    strs = file.read()
str_list = strs.split('\n')
numbers = ast.literal_eval(str_list[0])
numbers_rates = [round(x / sum(numbers) * 100) for x in numbers]
rates = ast.literal_eval(str_list[1])
goal_list = ast.literal_eval(str_list[2])
goal_rates = ast.literal_eval(str_list[3])


#填入人数行
for col_index,number in enumerate(numbers,start=1):
    table.cell(1,col_index).text = str(number)
#填入百分比行
for col_index,number_rate in enumerate(numbers_rates,start=1):
    table.cell(2,col_index).text = str(number_rate) + '%'

#填入题型得分率
table_question_rate = table.cell(4,2).tables[0]#选中题型得分率表格
for col_index in range(1,5):
    table_question_rate.cell(1,col_index).text = str(rates[col_index - 1])

#填入课程目标达成度情况表
table5 = document.tables[5]#第六个表格：课程目标达成度情况表
for row in range(1,4):
    for col_index in range(1,4):
        if col_index == 2:
            table5.cell(row,col_index).text = str(round(goal_list[row - 1][col_index - 1],2))
        else:
            table5.cell(row,col_index).text = str(goal_list[row - 1][col_index - 1])

#填入课程目标达成情况（分段统计）表
table5 = document.tables[6]#第七个表格：课程目标达成情况（分段统计）表
for col_index in range(1,7):
    for row in range(1,4):
        if col_index in range(2,7):
            table5.cell(row,col_index).text = str(goal_rates[6 - col_index][row - 1] * 100) + '%'
        else:
            #计算每一行的总达成度
            sum_rates = 0
            for i in range(2,7):
                sum_rates += float(table5.cell(row,i).text[:-1])
            table5.cell(row,col_index).text = str(sum_rates / 5) +'%'

#填入语句
table6 = document.tables[7]
with open(r"F:\代码\Python代码\课程达成度平台开发\pythonProject2\user_text_data.txt",'r',encoding='utf-8') as file:
    appraise_words = ''
    pattern = r'\[.*?\] \[.*?\] (.*)'  # 匹配 "[时间戳] [图片名] 评价内容"
    for line in file.readlines():
        if not line:
            continue
        match = re.match(pattern,line)
        if match:
            appraise_words += match.group(1)
table6.cell(5,1).text = appraise_words

document.save(r"C:\Users\lenovo\Desktop\分析报告.docx")

